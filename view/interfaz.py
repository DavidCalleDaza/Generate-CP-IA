import os
import subprocess
import webbrowser  # Importar para abrir el archivo de Excel
from docx import Document
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.popup import Popup



class CopiarTablaHU(BoxLayout):
    def __init__(self, **kwargs):
        super(CopiarTablaHU, self).__init__(**kwargs)
        self.orientation = 'vertical'
        
        self.label = Label(text='Ingresa el nombre de la historia de usuario (sin extensión):')
        self.add_widget(self.label)
        
        self.nombre_hu = TextInput(multiline=False)
        self.add_widget(self.nombre_hu)
        
        # Cambiar el texto del botón a "Extraer Criterios"
        self.boton_extraer_criterios = Button(text='Extraer Criterios')
        self.boton_extraer_criterios.bind(on_press=self.copiar_tabla_hu_y_pegar)
        self.add_widget(self.boton_extraer_criterios)

        # Botón para abrir el archivo de Excel con el color #86E3CE
        self.boton_abrir_excel = Button(text='Abrir Excel', background_color=(134/255, 227/255, 206/255, 1))  # Color #86E3CE
        self.boton_abrir_excel.bind(on_press=self.abrir_excel)
        self.add_widget(self.boton_abrir_excel)

        # Botón para borrar registro de color #FA897B
        self.boton_borrar_registro = Button(text='Borrar Registro', background_color=(250/255, 137/255, 123/255, 1))  # Color #FA897B
        self.boton_borrar_registro.bind(on_press=self.borrar_registro)
        self.add_widget(self.boton_borrar_registro)

    def copiar_tabla_hu_y_pegar(self, instance):
        nombre_hu = self.nombre_hu.text.strip()
        ruta_hu = os.path.join(r"C:\Users\d4vid\OneDrive\Escritorio\Generate-CP-IA\Archivos_imp\Historias de usuario", f"{nombre_hu}.docx")
        ruta_test = r"C:\Users\d4vid\OneDrive\Escritorio\Generate-CP-IA\Archivos_imp\Criterios_de_aceptacion.docx"

        try:
            doc_hu = Document(ruta_hu)
            print(f"Archivo {nombre_hu}.docx cargado correctamente.")
        except Exception as e:
            self.mostrar_popup(f"Error al cargar {nombre_hu}.docx: {e}")
            return
        
        try:
            doc_test = Document(ruta_test)
            print("Archivo test.docx cargado correctamente.")
        except Exception as e:
            self.mostrar_popup(f"Error al cargar test.docx: {e}")
            return

        tabla_copiar = None
        for table in doc_hu.tables:
            if "Criterios de Aceptación" in table.cell(0, 0).text:
                tabla_copiar = table
                break

        if not tabla_copiar:
            self.mostrar_popup("No se encontró la tabla 'Criterios de Aceptación'.")
            return

        for row in tabla_copiar.rows:
            nueva_fila = doc_test.add_table(rows=1, cols=len(row.cells)).rows[-1]
            for idx, cell in enumerate(row.cells):
                nueva_fila.cells[idx].text = cell.text

        try:
            doc_test.save(ruta_test)
            self.mostrar_popup("La tabla se ha copiado y test.docx ha sido guardado correctamente.")
        except Exception as e:
            self.mostrar_popup(f"Error al guardar test.docx: {e}")

        try:
            subprocess.run(["python", r"C:\Users\d4vid\OneDrive\Escritorio\Generate-CP-IA\model\hu.py"], check=True)
            print("El script hu.py se ha ejecutado correctamente.")
        except Exception as e:
            self.mostrar_popup(f"Error al ejecutar hu.py: {e}")

    def abrir_excel(self, instance):
        ruta_excel = r"C:\Users\d4vid\OneDrive\Escritorio\Generate-CP-IA\Archivos_xpo\VBA\CP-VBA.xlsm"
        try:
            webbrowser.open(ruta_excel)
            print("El archivo de Excel se ha abierto correctamente.")
        except Exception as e:
            self.mostrar_popup(f"Error al abrir el archivo de Excel: {e}")

    def borrar_registro(self, instance):
        try:
            subprocess.run(["python", r"C:\Users\d4vid\OneDrive\Escritorio\Generate-CP-IA\model\limpiar.py"], check=True)
            self.mostrar_popup("El script limpiar.py se ha ejecutado correctamente.")
        except Exception as e:
            self.mostrar_popup(f"Error al ejecutar limpiar.py: {e}")

    def mostrar_popup(self, mensaje):
        popup = Popup(title='Información', content=Label(text=mensaje), size_hint=(None, None), size=(400, 200))
        popup.open()


class MiApp(App):
    def build(self):
        return CopiarTablaHU()


if __name__ == '__main__':
    MiApp().run()