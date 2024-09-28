import os
from docx import Document
import subprocess  # Importar subprocess para ejecutar otro script

def copiar_tabla_hu_y_pegar():
    # Pedir el nombre de la historia de usuario
    nombre_hu = input("Ingresa el nombre de la historia de usuario (sin extensión, e.g., HU-003): ")
    
    # Rutas de los archivos
    ruta_hu = os.path.join(r"C:\Users\d4vid\OneDrive\Escritorio\Generate-CP-IA\Archivos_imp\Historias de usuario", f"{nombre_hu}.docx")
    ruta_test = r"C:\Users\d4vid\OneDrive\Escritorio\Generate-CP-IA\Archivos_imp\Criterios_de_aceptacion.docx"
    ruta_datos_generales = r"C:\Users\d4vid\OneDrive\Escritorio\Generate-CP-IA\Archivos_imp\Datos_generales_historia.docx"
    
    # Cargar el documento de la historia de usuario
    try:
        doc_hu = Document(ruta_hu)
        print(f"Archivo {nombre_hu}.docx cargado correctamente.")
    except Exception as e:
        print(f"Error al cargar {nombre_hu}.docx: {e}")
        return
    
    # Crear un nuevo documento para el archivo de criterios de aceptación si no existe
    try:
        doc_test = Document(ruta_test)
        print("Archivo Criterios_de_aceptacion.docx cargado correctamente.")
    except Exception as e:
        print(f"Error al cargar Criterios_de_aceptacion.docx: {e}")
        return

    # Crear un nuevo documento para datos generales
    doc_datos_generales = Document()

    # Buscar la tabla de "Criterios de Aceptación" en HU-003.docx
    tabla_copiar = None
    for table in doc_hu.tables:
        # Suponiendo que la tabla de "Criterios de Aceptación" es la primera
        if "Criterios de Aceptación" in table.cell(0, 0).text:  # Verificamos el encabezado de la tabla
            tabla_copiar = table
            break

    if not tabla_copiar:
        print("No se encontró la tabla 'Criterios de Aceptación'.")
        return

    # Copiar el contenido de la tabla de criterios de aceptación
    print("Copiando la tabla 'Criterios de Aceptación'...")
    for row in tabla_copiar.rows:
        nueva_fila = doc_test.add_table(rows=1, cols=len(row.cells)).rows[-1]
        for idx, cell in enumerate(row.cells):
            nueva_fila.cells[idx].text = cell.text

    # Guardar el archivo test.docx con la nueva tabla
    try:
        doc_test.save(ruta_test)
        print("La tabla se ha copiado y Criterios_de_aceptacion.docx ha sido guardado correctamente.")
    except Exception as e:
        print(f"Error al guardar Criterios_de_aceptacion.docx: {e}")

    # Extraer información de la primera tabla del documento de historia de usuario
    primera_tabla = doc_hu.tables[0]  # Suponiendo que la primera tabla es la que queremos
    print("Extrayendo información de la primera tabla...")
    
    for row in primera_tabla.rows:
        nueva_fila = doc_datos_generales.add_table(rows=1, cols=len(row.cells)).rows[-1]
        for idx, cell in enumerate(row.cells):
            nueva_fila.cells[idx].text = cell.text

    # Guardar el archivo Datos_generales_historia.docx con la información extraída
    try:
        doc_datos_generales.save(ruta_datos_generales)
        print("La información de la primera tabla se ha guardado en Datos_generales_historia.docx.")
    except Exception as e:
        print(f"Error al guardar Datos_generales_historia.docx: {e}")

    # Ejecutar el script hu.py después de terminar el proceso
    try:
        subprocess.run(["python", r"C:\Users\d4vid\OneDrive\Escritorio\Generate-CP-IA\model\hu.py"], check=True)
        print("El script hu.py se ha ejecutado correctamente.")
    except Exception as e:
        print(f"Error al ejecutar hu.py: {e}")

if __name__ == "__main__":
    copiar_tabla_hu_y_pegar()
